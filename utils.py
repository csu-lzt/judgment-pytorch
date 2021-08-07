# -*- coding:utf-8 -*-
import json, jieba, re, copy, docx, os, xlwt
from tqdm import tqdm
import numpy as np


##一些工具性函数，包括数据读取和分割等

def load_data(json_filename, return_length=False):
    """加载数据
    返回：[[text1各个句子组成的列表],[text2各个句子组成的列表]...]
    修改后：返回 [text各个句子组成的列表]，[对应标签的列表]
    """
    Sentence = []
    Label = []
    Length = []  # 每篇裁判文书有多少个句子
    with open(json_filename, 'r', encoding="utf8") as f:
        for line in tqdm(f, desc=u'读取数据中'):
            data = json.loads(line)
            single_text = data.get('text')
            # sentences_in_onetext = []
            length = 0
            for i, item in enumerate(single_text):
                sentence = item['sentence'].strip().replace('\u3000', '')
                if re.search("&#xD", sentence) or len(sentence) < 4 or re.compile(
                        r'^[-+]?[-0-9]\d*\.\d*|[-+]?\.?[0-9]\d*$').match(sentence):  ##基于正则规则，过滤一些杂乱的句子
                    continue  # 处理后最大句子数量为244，<256
                # sentences_in_onetext.append(sentence)
                length += 1
                Sentence.append(sentence)
                Label.append(item['label'])
            Length.append(length)
            # print('len of text list', len(texts),texts)  作者自己分的句，所以句子多，len(texts)长
            # D.append(sentences_in_onetext)
    if return_length == True:
        return Sentence, Label, Length
    return Sentence, Label


def split_data(raw_data, out_path, write_mode='a', train_num=11530, valid_num=1000, test_num=1000):
    """随机切分训练集、验证集、测试集
    raw_data: 原始数据集路径
    out_path: 输出路径文件夹，末尾不用加/
    write_mode : a在原来的基础上继续写 r清空原有的重新写入
    """
    train_data, valid_data, test_data = out_path + '/train_data.json', out_path + '/valid_data.json', out_path + '/test_data.json'
    count = 0
    with open(raw_data, 'r', encoding='utf8') as raw_f:
        for line in tqdm(raw_f, desc='读取原始数据中'):
            count += 1
    print('文件总个数(行数):', count)
    # 随机打乱索引，划分数据集
    index = np.arange(count)
    np.random.seed(2021)  # 固定种子
    np.random.shuffle(index)
    train_index, valid_index, test_index = index[:train_num], index[train_num: train_num + valid_num], index[
                                                                                                       train_num + valid_num:train_num + valid_num + test_num]
    with open(raw_data, 'r', encoding='utf8') as raw_f:
        for i, line in enumerate(tqdm(raw_f)):
            if i in train_index:
                with open(train_data, write_mode, encoding='utf8') as train_f:
                    train_f.write(line)
            elif i in valid_index:
                with open(valid_data, write_mode, encoding='utf8') as train_f:
                    train_f.write(line)
            elif i in test_index:
                with open(test_data, write_mode, encoding='utf8') as train_f:
                    train_f.write(line)
    return 0


# 这里有问题，是CAIL阶段遗留下的且未被发现的，注意不要犯错误（把逗号等符号去掉了再分词）
def str2list(str):
    # punctuation = """。：，、《》（）"""
    # re_punctuation = "[{}]+".format(punctuation)
    # str = str.replace(" ", "")
    # str = re.sub(re_punctuation, "", str)
    # 正则式去除括号内的内容
    re_str = re.sub(r'[\(|（|【].*?[\)|）|】]', "", str)
    # 一个标点符号表，分词后去除句子中的标点符号
    remove_list = [' ', ',', '.', '，', '。', ':', '：', '《', '》', '(', ')', '（', '）', '、']
    grams_remove = []
    grams = jieba.lcut(re_str)  # 直接切分成列表
    for gram in grams:
        if gram not in remove_list:
            grams_remove.append(gram)
    return grams_remove


def Rouge_1(grams_model, grams_reference):  # terms_reference为参考摘要，terms_model为候选摘要   ***one-gram*** 一元模型
    temp_precision = 0  # 准确率（precision） （你给出的结果有多少是正确的）
    temp_recall = 0  # 召回率（recall） （正确的结果有多少被你给出了）
    grams_reference_all = len(grams_reference)
    grams_model_all = len(grams_model)
    for i in range(grams_reference_all):
        c = grams_reference[i - temp_recall]  # remove了多少个出去，就要往前移多少个，确保下标不会出错
        if c in grams_model:
            grams_reference.remove(c)
            grams_model.remove(c)
            temp_recall = temp_recall + 1
    temp_precision = temp_recall
    precision = temp_precision / grams_model_all
    recall = temp_recall / grams_reference_all
    if temp_recall == 0:
        Fscore = 0
    else:
        Fscore = (2 * precision * recall) / (precision + recall)
    # print(u'准确率：',precision)
    # print(u'召回率：',recall)
    # print(u'R1：', Fscore)

    return [Fscore, precision, recall]


def Rouge_2(grams_model, grams_reference):  # terms_reference为参考摘要，terms_model为候选摘要   ***Bi-gram***  2元模型
    temp_precision = 0  # 准确率（precision） （你给出的结果有多少是正确的）
    temp_recall = 0  # 召回率（recall） （正确的结果有多少被你给出了）
    grams_reference_all = len(grams_reference) - 1  # 这里减1代表2元组的个数
    grams_model_all = len(grams_model) - 1
    gram_2_model = []
    gram_2_reference = []
    for x in range(grams_model_all):
        gram_2_model.append(grams_model[x] + grams_model[x + 1])
    for x in range(grams_reference_all):
        gram_2_reference.append(grams_reference[x] + grams_reference[x + 1])

    for i in range(grams_reference_all):
        c = gram_2_reference[i - temp_recall]
        if c in gram_2_model:
            gram_2_reference.remove(c)
            gram_2_model.remove(c)
            temp_recall = temp_recall + 1
    temp_precision = temp_recall
    precision = temp_precision / grams_model_all
    recall = temp_recall / grams_reference_all
    if temp_recall == 0:
        Fscore = 0
    else:
        Fscore = (2 * precision * recall) / (precision + recall)
    #     print(u'准确率：',precision)
    #     print(u'召回率：',recall)
    # print(u'R2：', Fscore)
    return [Fscore, precision, recall]


def LCS(string1, string2):
    len1 = len(string1)
    len2 = len(string2)
    res = [[0 for i in range(len1 + 1)] for j in range(len2 + 1)]
    for i in range(1, len2 + 1):
        for j in range(1, len1 + 1):
            if string2[i - 1] == string1[j - 1]:
                res[i][j] = res[i - 1][j - 1] + 1
            else:
                res[i][j] = max(res[i - 1][j], res[i][j - 1])
    return res[-1][-1]


def Rouge_L(grams_model, grams_reference):  # terms_reference为参考摘要，terms_model为候选摘要   ***one-gram*** 一元模型
    grams_reference_all = len(grams_reference)
    grams_model_all = len(grams_model)
    LCS_n = LCS(grams_model, grams_reference)
    precision = LCS_n / grams_model_all
    recall = LCS_n / grams_reference_all
    if recall == 0:
        Fscore = 0
    else:
        Fscore = (2 * precision * recall) / (precision + recall)
    #     print(u'准确率：',precision)
    #     print(u'召回率：',recall)
    # print(u'RL：', Fscore)
    # print(u'最长子序列', LCS_n)
    return [Fscore, precision, recall]


def Rouge(grams_model, grams_reference):
    grams_model1 = copy.deepcopy(grams_model)
    grams_model2 = copy.deepcopy(grams_model)
    grams_model3 = copy.deepcopy(grams_model)
    grams_reference1 = copy.deepcopy(grams_reference)
    grams_reference2 = copy.deepcopy(grams_reference)
    grams_reference3 = copy.deepcopy(grams_reference)
    rouge_1_F1, rouge_1_precision, rouge_1_recall = Rouge_1(grams_model1, grams_reference1)
    rouge_2_F1, rouge_2_precision, rouge_2_recall = Rouge_2(grams_model2, grams_reference2)
    rouge_L_F1, rouge_L_precision, rouge_L_recall = Rouge_L(grams_model3, grams_reference3)
    rouge_all_F1 = 0.2 * rouge_1_F1 + 0.4 * rouge_2_F1 + 0.4 * rouge_L_F1
    rouge_all_precison = 0.2 * rouge_1_precision + 0.4 * rouge_2_precision + 0.4 * rouge_L_precision
    rouge_all_recall = 0.2 * rouge_1_recall + 0.4 * rouge_2_recall + 0.4 * rouge_L_recall

    rouge_1 = [rouge_1_F1, rouge_1_precision, rouge_1_recall]
    rouge_2 = [rouge_2_F1, rouge_2_precision, rouge_2_recall]
    rouge_L = [rouge_L_F1, rouge_L_precision, rouge_L_recall]
    rouge_all = [rouge_all_F1, rouge_all_precison, rouge_all_recall]
    rouge = np.array([rouge_1, rouge_2, rouge_L, rouge_all])
    return rouge


def evaluate_rouge(predict_summary, golden_summary):
    ROUGE = np.zeros((4, 3))
    predict_summary_grams = str2list(predict_summary)
    golden_summary_grams = str2list(golden_summary)
    if predict_summary != '':
        ROUGE = 100 * Rouge(predict_summary_grams, golden_summary_grams)
    # rouge1_F1, rouge1_P, rouge1_R,
    # rouge2_F1, rouge2_P, rouge2_R,
    # rougeL_F1, rougeL_P, rougeL_R,
    # rouge_F1, rouge_P, rouge_R
    # 返回一个4*3的np矩阵
    return ROUGE


def del_file(path):
    for i in os.listdir(path):  # os.listdir(path_data)#返回一个列表，里面是当前目录下面的所有东西的相对路径
        file_data = path + "\\" + i  # 当前文件夹的下面的所有东西的绝对路径
        if os.path.isfile(file_data) == True:  # os.path.isfile判断是否为文件,如果是文件,就删除.如果是文件夹.递归给del_file.
            os.remove(file_data)
        else:
            del_file(file_data)


def write_docx(sentences, label1, pre_label1, gold, pred, i, path):
    '''
    :param sentences: 裁判文书句子列表
    :param label1: 真实关键句列表
    :param pre_label1: 预测关键句列表
    :param gold: 人工摘要
    :param pred: 生成摘要
    :param i: 文书序号
    :param path: 写入的文件夹路径（相对），最后不加/
    :return: 无，写入各个word文档,并用色彩区分不同的句子类型
    '''
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    file = docx.Document()
    file.styles['Normal'].font.name = u'宋体'
    file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    p1 = file.add_paragraph('原文：')
    # add_run在同一段添加内容
    for sentence in sentences:
        if sentence in label1 and sentence not in pre_label1:
            p1.add_run(sentence).font.color.rgb = RGBColor(255, 0, 0)  # 真实有，预测没有，红色
        elif sentence not in label1 and sentence in pre_label1:
            p1.add_run(sentence).font.color.rgb = RGBColor(0, 0, 255)  # 真实没有，预测有，蓝色
        elif sentence in label1 and sentence in pre_label1:
            p1.add_run(sentence).font.color.rgb = RGBColor(255, 0, 255)  # 真实和预测都有，紫色
        else:
            p1.add_run(sentence)
    file.add_paragraph('人工摘要：' + gold)
    file.add_paragraph('生成摘要：' + pred)
    file.save(path + f'/{i}.docx')


def select_summary(data_path='data/select_76_data/data_76.json', pred_path='data/select_76_data/pre_summary_76.json',
                   label1_path='data/select_76_data/gold_summary_label1.json', high_socre=65):
    # 调查问卷选取好一点的司法摘要
    predict_summary = []
    golden_summary = []
    sentences_of_all = []
    label1s = []
    pre_label1s = []
    with open(data_path, 'r', encoding="utf8") as f:
        for line in tqdm(f, desc=u'读取原始数据中'):
            single_judgment = []
            data = json.loads(line)
            golden_summary.append(data.get('summary'))
            text = data.get('text')
            for i, item in enumerate(text):
                sentence = item['sentence'].strip().replace('\u3000', '')
                single_judgment.append(sentence)
            sentences_of_all.append(single_judgment)
    with open(pred_path, 'r', encoding='utf8') as fp:
        for line in tqdm(fp, desc='读取预测的摘要'):
            data = json.loads(line)
            predict_summary.append(data.get('summary'))
    with open(label1_path, 'r', encoding='utf8') as fl:
        for line in tqdm(fl, desc='读取关键句中'):
            data = json.loads(line)
            label1s.append(data.get('label1'))
            pre_label1s.append(data.get('pre_label1'))

    rouge_all = 0
    for i, pred in enumerate(predict_summary):
        # pred=predict_summary[i]
        gold = golden_summary[i]
        sentences = sentences_of_all[i]
        label1 = label1s[i]
        pre_label1 = pre_label1s[i]
        ROUGE = evaluate_rouge(pred, gold)  # rouge分值矩阵
        rouge = ROUGE[3][0]  # rouge总分值F1
        print(rouge)
        rouge_all += rouge
        if rouge > high_socre:
            file_path = 'word_judgment/higher'
            write_docx(sentences, label1, pre_label1, gold, pred, i, file_path)
        elif rouge > 50 and rouge < 51:
            file_path = 'word_judgment/middle'
            write_docx(sentences, label1, pre_label1, gold, pred, i, file_path)
        elif rouge < 35:
            file_path = 'word_judgment/lower'
            write_docx(sentences, label1, pre_label1, gold, pred, i, file_path)
            print(pred, gold)
    print(rouge_all / 76)


def write_excel(index, data_path='data/data_all.json'):
    # 创建excel工作表
    workbook = xlwt.Workbook(encoding='utf-8')
    worksheet = workbook.add_sheet('sheet1')
    # 设置表头
    worksheet.write(0, 0, label='序号')
    worksheet.write(0, 1, label='ID')
    worksheet.write(0, 2, label='参考摘要')
    worksheet.write(0, 3, label='句子')
    worksheet.write(0, 4, label='参考标签')
    worksheet.write(0, 5, label='标签A')
    worksheet.write(0, 6, label='标签B')
    # 设置表格每列宽度
    worksheet.col(0).width = 1000
    worksheet.col(1).width = 1000
    worksheet.col(2).width = 14500
    worksheet.col(3).width = 20000
    worksheet.col(4).width = 2000
    worksheet.col(5).width = 2000
    worksheet.col(6).width = 2000
    # 读文件并写入excel
    with open(data_path, 'r', encoding="utf8") as f:
        start_book_id = (index - 1) * 10
        end_book_id = (index - 1) * 10 + 1
        book_id = 0  # 文书序号
        raw_id = 1  # 行号
        for line in tqdm(f):
            book_id += 1
            if book_id < start_book_id:
                continue
            elif book_id > end_book_id:
                break
            sentence_list = []
            label_list = []
            data = json.loads(line)
            id = data.get('id')
            summary = data.get('summary')
            text = data.get('text')
            for item in text:
                sentence = item['sentence'].strip().replace('\u3000', '')
                if re.search("&#xD", sentence) or len(sentence) < 4 or re.compile(
                        r'^[-+]?[-0-9]\d*\.\d*|[-+]?\.?[0-9]\d*$').match(sentence):  ##基于正则规则，过滤一些杂乱的句子
                    continue
                sentence_list.append(sentence)
                label_list.append(item["label"])
            n = len(label_list)  # 一篇文书有n句话
            # 合并单元格
            # table.write_merge(x, x + m, y, y + n, string, style)
            # x表示行，y表示列，m表示跨行个数，n表示跨列个数，string表示要写入的单元格内容，style表示单元格样式。
            worksheet.write_merge(raw_id, raw_id + n - 1, 0, 0, book_id)
            worksheet.write_merge(raw_id, raw_id + n - 1, 1, 1, id)
            worksheet.write_merge(raw_id, raw_id + n - 1, 2, 2, summary)
            for i in range(n):
                worksheet.write(raw_id + i, 3, sentence_list[i])
                worksheet.write(raw_id + i, 4, label_list[i])
            raw_id += n
    save_path = f'excel_judgment/{index}.xlsx'
    workbook.save(save_path)


if __name__ == '__main__':
    # del_file('word_judgment/higher')
    # del_file('word_judgment/middle')
    # del_file('word_judgment/lower')
    # select_summary()
    # write_excel(6)
    # for i in range(5):
    #     write_excel(i+1)
    _, _, length = load_data('data/classify_data/train_data.json', return_length=True)
    print(length)